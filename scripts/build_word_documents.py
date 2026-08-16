#!/usr/bin/env python3
"""Build polished Word versions of the Una Voce proposal and prospectus."""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "word"
SKILL_ROOT = Path(
    "/Users/richcuff/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.805.11740/skills/documents"
)

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "18324A"
GRAY = "5F6872"
LIGHT_GRAY = "F4F6F9"
BORDER = "CDD4DC"
WHITE = "FFFFFF"
BLACK = "000000"


@dataclass(frozen=True)
class Preset:
    name: str
    body_after: float
    body_line: float
    body_alignment: WD_ALIGN_PARAGRAPH
    h1_before: float
    h1_after: float
    h2_before: float
    h2_after: float
    h3_before: float
    h3_after: float
    list_after: float
    list_line: float
    list_left_dxa: int
    list_hanging_dxa: int
    table_header_fill: str


GRANT_PRESET = Preset(
    name="grant_proposal",
    body_after=6,
    body_line=1.25,
    body_alignment=WD_ALIGN_PARAGRAPH.LEFT,
    h1_before=16,
    h1_after=8,
    h2_before=12,
    h2_after=6,
    h3_before=8,
    h3_after=4,
    list_after=4,
    list_line=1.208,
    list_left_dxa=540,
    list_hanging_dxa=280,
    table_header_fill=LIGHT_GRAY,
)

NARRATIVE_PRESET = Preset(
    name="narrative_proposal",
    body_after=8,
    body_line=1.333,
    body_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    h1_before=18,
    h1_after=10,
    h2_before=12,
    h2_after=6,
    h3_before=8,
    h3_after=4,
    list_after=4,
    list_line=1.208,
    list_left_dxa=540,
    list_hanging_dxa=280,
    table_header_fill=LIGHT_GRAY,
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=BLACK, bold=None, italic=None) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def configure_styles(doc: Document, preset: Preset) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(preset.body_after)
    normal.paragraph_format.line_spacing = preset.body_line
    normal.paragraph_format.alignment = preset.body_alignment
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, BLUE, bold=True)
    # Named renderer-clearance override: LibreOffice needs slightly more leading
    # and trailing space than Word to avoid glyph collisions at page boundaries.
    h1.paragraph_format.space_before = Pt(max(preset.h1_before, 18))
    h1.paragraph_format.space_after = Pt(max(preset.h1_after, 12))
    h1.paragraph_format.line_spacing = 1.15
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(max(preset.h2_before, 14))
    h2.paragraph_format.space_after = Pt(max(preset.h2_after, 10))
    h2.paragraph_format.line_spacing = 1.15
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(max(preset.h3_before, 10))
    h3.paragraph_format.space_after = Pt(max(preset.h3_after, 8))
    h3.paragraph_format.line_spacing = 1.15
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    styles = doc.styles
    for style_name in ("UV Bullet", "UV Number"):
        if style_name in styles:
            style = styles[style_name]
        else:
            style = styles.add_style(style_name, 1)
        set_style_font(style, 11, BLACK)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(max(preset.list_after, 6))
        style.paragraph_format.line_spacing = max(preset.list_line, 1.25)
        style.paragraph_format.widow_control = True

    if "UV Callout" in styles:
        callout = styles["UV Callout"]
    else:
        callout = styles.add_style("UV Callout", 1)
    set_style_font(callout, 9.5, GRAY, italic=True)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.15

    if "UV Table" in styles:
        table_style = styles["UV Table"]
    else:
        table_style = styles.add_style("UV Table", 1)
    set_style_font(table_style, 9.25, BLACK)
    table_style.paragraph_format.space_before = Pt(0)
    table_style.paragraph_format.space_after = Pt(0)
    table_style.paragraph_format.line_spacing = 1.12

    if "UV Table Header" in styles:
        table_header = styles["UV Table Header"]
    else:
        table_header = styles.add_style("UV Table Header", 1)
    set_style_font(table_header, 9.25, NAVY, bold=True)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.05


def add_abstract_numbering(doc: Document, preset: Preset, fmt: str) -> int:
    numbering = doc.part.numbering_part.element
    existing = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    list_left_dxa = 720 if fmt == "decimal" else preset.list_left_dxa
    list_hanging_dxa = 360 if fmt == "decimal" else preset.list_hanging_dxa
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(list_left_dxa))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(list_left_dxa))
    ind.set(qn("w:hanging"), str(list_hanging_dxa))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(int(max(preset.list_after, 6) * 20)))
    spacing.set(qn("w:line"), str(round(max(preset.list_line, 1.25) * 240)))
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)

    if fmt == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)

    abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def instantiate_numbering(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    existing = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    # Explicit restart is required for LibreOffice; a fresh numId alone may
    # continue the most recent list based on the same abstract definition.
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_inline_runs(paragraph, text: str, *, base_size=None, base_color=None, base_bold=False) -> None:
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size, color=base_color, bold=base_bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, color=base_color, bold=base_bold, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size or 9.5, color=base_color or DARK_BLUE, bold=False, name="Courier New")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color, bold=base_bold)


def paragraph_border_and_shading(paragraph, fill=LIGHT_GRAY, border_color=BLUE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=9, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def configure_cover_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    # Keep the cover and body in one Word section. This avoids office-suite
    # page-style inheritance differences at a section boundary while still
    # allowing a clean cover through the first-page header/footer parts.
    doc.settings.odd_and_even_pages_header_footer = True

    for header in (section.header, section.even_page_header, section.first_page_header):
        header.is_linked_to_previous = False
        header.paragraphs[0].text = ""
    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = ""

    pff = section.first_page_footer.paragraphs[0]
    pff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pff.paragraph_format.space_after = Pt(0)
    run = pff.add_run("WORKING DRAFT  |  AUGUST 2026")
    set_run_font(run, size=8.5, color=GRAY, bold=True)

    # Number the hidden cover as zero, so the first visible body page is 1.
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), "0")


def start_content_section(doc: Document, running_label: str) -> None:
    doc.add_page_break()
    section = doc.sections[0]

    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(running_label.upper())
        set_run_font(run, size=8.5, color=GRAY, bold=True)

    for footer in (section.footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.paragraph_format.space_after = Pt(0)
        add_page_field(fp)


def add_spacer(doc: Document, points: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    p.add_run("")


def add_cover_ccc(doc: Document) -> None:
    add_spacer(doc, 32)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("UNA VOCE")
    set_run_font(run, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("A National Digital Gateway\nto the Liturgy of the Hours")
    set_run_font(run, size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Catholic Communication Campaign Project Proposal")
    set_run_font(run, size=14, color=GRAY, bold=True)

    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680], indent_dxa=120)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Word's accessibility model expects a declared header row for data
    # tables. Repeating it has no visual effect on this two-row cover summary.
    set_repeat_table_header(table.rows[0])
    labels = [
        ("AMOUNT REQUESTED", "$60,000"),
        ("PROJECT PERIOD", "12 months"),
        ("SCOPE", "National web and mobile"),
        ("ACCESS MODEL", "Free core experience"),
    ]
    for idx, (label, value) in enumerate(labels):
        cell = table.rows[idx // 2].cells[idx % 2]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(label)
        set_run_font(r1, size=8.5, color=BLUE, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=11, color=NAVY, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_borders(table, color=WHITE, size="12")

    add_spacer(doc, 22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Discovery  •  Habit  •  Community  •  Prayer")
    set_run_font(run, size=10.5, color=GRAY, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Working proposal for board, adviser, and grant-review discussion")
    set_run_font(run, size=9, color=GRAY)


def add_cover_prospectus(doc: Document) -> None:
    add_spacer(doc, 105)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("PROSPECTUS")
    set_run_font(run, size=10.5, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("Una Voce")
    set_run_font(run, size=32, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Discovery, Habit, and the Church's Daily Prayer")
    set_run_font(run, size=15, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(54)
    run = p.add_run("Business Plan and Possible Roadmaps")
    set_run_font(run, size=11.5, color=GRAY, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    p.paragraph_format.space_after = Pt(42)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(
        "Helping Catholics find a way into the Liturgy of the Hours, "
        "return tomorrow, and remain connected to the ministries already praying and producing it."
    )
    set_run_font(run, size=12, color=GRAY, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("WORKING DISCUSSION DOCUMENT")
    set_run_font(run, size=9, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("August 2026")
    set_run_font(run, size=10.5, color=GRAY)


def split_table_row(line: str) -> list[str]:
    content = line.strip().strip("|")
    return [cell.strip() for cell in content.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def choose_table_widths(headers: list[str]) -> list[int]:
    cols = len(headers)
    lower = [re.sub(r"[*`]", "", h).lower() for h in headers]
    if cols == 2:
        if any("amount" in h for h in lower):
            return [7200, 2160]
        if lower[0] in {"field", "stage"}:
            return [2500, 6860]
        return [3100, 6260]
    if cols == 3:
        if "maximum hours" in lower or "maximum amount" in lower:
            return [5600, 1600, 2160]
        if lower[:2] == ["period", "primary work"]:
            return [1700, 3400, 4260]
        if lower[0] == "stage" and "possible funding profile" in lower:
            return [1800, 3200, 4360]
        if lower[0] == "risk":
            return [1800, 3200, 4360]
        if lower[0] == "stage" and "una voce contribution" in lower:
            return [1700, 3500, 4160]
        return [2500, 3200, 3660]
    if cols == 4:
        return [1400, 2700, 2630, 2630]
    base = CONTENT_WIDTH_DXA // cols
    widths = [base] * cols
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]], preset: Preset) -> None:
    headers = rows[0]
    widths = choose_table_widths(headers)
    table = doc.add_table(rows=len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    numeric_headers = {"amount", "maximum hours", "maximum amount", "request"}
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        for c_idx, raw_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                set_cell_shading(cell, preset.table_header_fill)
            p = cell.paragraphs[0]
            p.style = "UV Table Header" if r_idx == 0 else "UV Table"
            p.paragraph_format.keep_together = True
            header_name = re.sub(r"[*`]", "", headers[c_idx]).strip().lower()
            if header_name in numeric_headers or (r_idx > 0 and re.fullmatch(r"\$?[\d,]+(?:\.\d+)?", re.sub(r"[*]", "", raw_text))):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif r_idx == 0 and len(raw_text) < 22:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(
                p,
                raw_text,
                base_size=9.25,
                base_color=NAVY if r_idx == 0 else BLACK,
                base_bold=(r_idx == 0),
            )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def add_markdown_body(
    doc: Document,
    markdown_path: Path,
    preset: Preset,
    page_break_headings: set[str] | None = None,
) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    bullet_abstract = add_abstract_numbering(doc, preset, "bullet")
    decimal_abstract = add_abstract_numbering(doc, preset, "decimal")
    bullet_num_id = instantiate_numbering(doc, bullet_abstract)
    decimal_num_id = None
    previous_kind = None

    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            previous_kind = None
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            heading_text = heading_match.group(2)
            if page_break_headings and heading_text in page_break_headings:
                p.paragraph_format.page_break_before = True
            p.paragraph_format.space_before = Pt(
                {1: max(preset.h1_before, 18), 2: max(preset.h2_before, 14), 3: max(preset.h3_before, 10)}[level]
            )
            p.paragraph_format.space_after = Pt(
                {1: max(preset.h1_after, 12), 2: max(preset.h2_after, 10), 3: max(preset.h3_after, 8)}[level]
            )
            p.paragraph_format.line_spacing = 1.15
            add_inline_runs(p, heading_text)
            previous_kind = "heading"
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            p = doc.add_paragraph(style="UV Callout")
            paragraph_border_and_shading(p)
            add_inline_runs(p, " ".join(quote_lines), base_size=9.5, base_color=GRAY)
            previous_kind = "callout"
            continue

        if "|" in stripped and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = split_table_row(lines[i])
                if len(row) == len(rows[0]):
                    rows.append(row)
                i += 1
            add_markdown_table(doc, rows, preset)
            previous_kind = "table"
            continue

        bullet_match = re.match(r"^\s*-\s+(.+)$", raw)
        if bullet_match:
            p = doc.add_paragraph(style="UV Bullet")
            apply_numbering(p, bullet_num_id)
            add_inline_runs(p, bullet_match.group(1))
            previous_kind = "bullet"
            i += 1
            continue

        number_match = re.match(r"^\s*\d+\.\s+(.+)$", raw)
        if number_match:
            if previous_kind != "number":
                decimal_num_id = instantiate_numbering(doc, decimal_abstract)
            p = doc.add_paragraph(style="UV Number")
            apply_numbering(p, decimal_num_id)
            add_inline_runs(p, number_match.group(1))
            previous_kind = "number"
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate:
                break
            if re.match(r"^(#{1,4})\s+", candidate):
                break
            if candidate.startswith(">") or re.match(r"^\s*-\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i]):
                break
            if "|" in candidate and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                break
            paragraph_lines.append(candidate)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(paragraph_lines))
        previous_kind = "paragraph"


def add_document_properties(doc: Document, title: str, subject: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Una Voce"
    props.keywords = "Una Voce, Liturgy of the Hours, Catholic, nonprofit"
    props.comments = "Working draft generated from the approved Markdown source."
    props.created = datetime(2026, 8, 6)
    props.modified = datetime(2026, 8, 6)


def audit_document(doc: Document, preset: Preset) -> None:
    for section in doc.sections:
        assert section.page_width.twips == PAGE_WIDTH_DXA
        assert section.page_height.twips == PAGE_HEIGHT_DXA
        assert section.left_margin.twips == 1440
        assert section.right_margin.twips == 1440
        assert section.top_margin.twips == 1440
        assert section.bottom_margin.twips == 1440
    assert abs(doc.styles["Normal"].paragraph_format.line_spacing - preset.body_line) < 0.001
    assert doc.styles["Normal"].font.name == "Calibri"
    for table in doc.tables:
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA, grid_widths
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == CONTENT_WIDTH_DXA
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == TABLE_INDENT_DXA


def build_document(source: Path, output: Path, preset: Preset, kind: str) -> None:
    doc = Document()
    configure_styles(doc, preset)
    configure_cover_page(doc)
    if kind == "ccc":
        add_cover_ccc(doc)
        start_content_section(doc, "Una Voce  |  CCC Project Proposal")
        title = "Una Voce - $60,000 CCC Project Proposal"
        subject = "Catholic Communication Campaign project proposal"
        page_break_headings = {"Project thesis"}
    else:
        add_cover_prospectus(doc)
        start_content_section(doc, "Una Voce  |  Prospectus and Roadmaps")
        title = "Una Voce - Prospectus and Possible Roadmaps"
        subject = "Generalized business plan and strategic roadmap"
        page_break_headings = None

    add_markdown_body(doc, source, preset, page_break_headings=page_break_headings)
    add_document_properties(doc, title, subject)
    audit_document(doc, preset)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_document(
        ROOT / "docs" / "ccc-60000-proposal.md",
        OUTPUT_DIR / "Una Voce - CCC 60000 Proposal.docx",
        GRANT_PRESET,
        "ccc",
    )
    build_document(
        ROOT / "docs" / "una-voce-prospectus.md",
        OUTPUT_DIR / "Una Voce - Prospectus and Roadmaps.docx",
        NARRATIVE_PRESET,
        "prospectus",
    )


if __name__ == "__main__":
    main()
